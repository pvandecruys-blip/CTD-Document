"""
DOCX generation engine for CTD stability sections.

Produces Word documents following CTD Module 3 formatting conventions,
including:
  - Title/header with CTD section numbering
  - Table of Contents
  - List of Tables
  - Narrative paragraph
  - Summary table
  - Detail tables per lot × condition

Uses python-docx for programmatic document construction.
"""

import json
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional
from uuid import uuid4

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


# ── Data structures for generation input ────────────────────────────

@dataclass
class LotInfo:
    lot_id: str
    lot_number: str
    manufacturer: Optional[str] = None
    manufacturing_site: Optional[str] = None
    intended_use: Optional[str] = None
    lot_use_label: Optional[str] = None


@dataclass
class ConditionInfo:
    condition_id: str
    label: str
    display_order: int = 0


@dataclass
class TimepointInfo:
    timepoint_id: str
    label: str
    sort_order: int = 0


@dataclass
class AttributeRow:
    """One row in a stability detail table."""
    attribute_name: str
    method_group: Optional[str]
    acceptance_criteria: Optional[str]
    values: dict[str, str]  # timepoint_label → value/status


@dataclass
class DetailTableData:
    """Data for one detail stability table (one lot × one condition)."""
    lot: LotInfo
    condition: ConditionInfo
    timepoints: list[TimepointInfo]
    rows: list[AttributeRow]
    footnotes: list[str] = field(default_factory=list)


@dataclass
class SummaryTableRow:
    lot_number: str
    lot_use: str
    study_start: Optional[str]
    conditions_presented: str
    detail_table_refs: str  # e.g., "Tables 3.2.P.8.3-2 to 3.2.P.8.3-5"


@dataclass
class GenerationInput:
    """Complete input for document generation."""
    product_name: str
    product_type: str  # "drug_substance" or "drug_product"
    ctd_section: str   # e.g., "3.2.P.8.3" or "3.2.S.7"
    section_title: str  # e.g., "Stability Data – Accelerated"
    study_type_label: str  # e.g., "Accelerated"

    lots: list[LotInfo]
    conditions: list[ConditionInfo]
    timepoints: list[TimepointInfo]
    detail_tables: list[DetailTableData]
    summary_rows: list[SummaryTableRow]

    # Styling
    confidentiality_mark: Optional[str] = None
    header_text: Optional[str] = None
    footer_text: Optional[str] = None

    # Footnotes
    global_footnotes: list[str] = field(default_factory=lambda: [
        "S = Meets acceptance criteria at the time of testing.",
        "NS = Does not meet acceptance criteria.",
        "NT = Not tested at this timepoint.",
        "W = week; M = month.",
        "Acceptance criteria in place at time of testing.",
    ])


@dataclass
class TraceabilityEntry:
    """One entry in the traceability appendix."""
    table_number: str
    row_attribute: str
    column_timepoint: str
    value: str
    source_document: str
    source_page: Optional[int]
    source_table_ref: Optional[str]
    source_snippet: Optional[str]


# ── DOCX Generator ─────────────────────────────────────────────────

class StabilityDocxGenerator:
    """
    Generates a CTD-compliant stability section DOCX document.

    Usage:
        gen = StabilityDocxGenerator(input_data)
        doc_path = gen.generate("output/stability.docx")
    """

    # Style constants
    FONT_NAME = "Times New Roman"
    HEADING_FONT = "Arial"
    TABLE_FONT_SIZE = Pt(8)
    BODY_FONT_SIZE = Pt(11)
    HEADER_BG_COLOR = "D9E2F3"  # Light blue header
    TABLE_BORDER_COLOR = "000000"

    def __init__(self, input_data: GenerationInput):
        self.data = input_data
        self.doc = Document()
        self.table_counter = 0
        self.traceability: list[TraceabilityEntry] = []

    def generate(self, output_path: str) -> str:
        """Generate the complete DOCX document and save to output_path."""
        self._setup_styles()
        self._add_header_footer()
        self._add_title()
        self._add_toc()
        self._add_list_of_tables()
        self._add_narrative()
        self._add_summary_table()
        self._add_detail_tables()

        # Save
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"DOCX generated: {output_path}")
        return output_path

    def generate_traceability_json(self, output_path: str) -> str:
        """Export traceability appendix as JSON."""
        data = [
            {
                "table_number": e.table_number,
                "row_attribute": e.row_attribute,
                "column_timepoint": e.column_timepoint,
                "value": e.value,
                "source_document": e.source_document,
                "source_page": e.source_page,
                "source_table_ref": e.source_table_ref,
                "source_snippet": e.source_snippet,
            }
            for e in self.traceability
        ]
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return output_path

    # ── Setup ───────────────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = self.BODY_FONT_SIZE

        # Configure heading styles
        for level in range(1, 5):
            style_name = f"Heading {level}"
            if style_name in self.doc.styles:
                heading_style = self.doc.styles[style_name]
                heading_style.font.name = self.HEADING_FONT
                heading_style.font.color.rgb = RGBColor(0, 0, 0)

        # Set narrow margins
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.0)

    def _add_header_footer(self):
        """Add header and footer to all pages."""
        section = self.doc.sections[0]

        # Header
        header = section.header
        header.is_linked_to_previous = False
        if self.data.confidentiality_mark:
            p = header.paragraphs[0]
            p.text = self.data.confidentiality_mark
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.style.font.size = Pt(9)
            p.style.font.italic = True

        if self.data.header_text:
            p = header.add_paragraph()
            p.text = self.data.header_text
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.style.font.size = Pt(9)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        if self.data.footer_text:
            p = footer.paragraphs[0]
            p.text = self.data.footer_text
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style.font.size = Pt(9)

    # ── Title ───────────────────────────────────────────────────────

    def _add_title(self):
        """Add the section title heading."""
        title = f"{self.data.ctd_section} {self.data.section_title}"
        self.doc.add_heading(title, level=1)

    # ── Table of Contents ───────────────────────────────────────────

    def _add_toc(self):
        """Add a Table of Contents field (auto-updated by Word)."""
        self.doc.add_heading("Table of Contents", level=2)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar)

        run = paragraph.add_run("(Table of contents will be updated when you open the document in Word)")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar)

        self.doc.add_paragraph()  # spacer

    # ── List of Tables ──────────────────────────────────────────────

    def _add_list_of_tables(self):
        """Add a List of Tables field."""
        self.doc.add_heading("List of Tables", level=2)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar)

        run = paragraph.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve">'
            f' TOC \\h \\z \\c "Table" </w:instrText>'
        )
        run._r.append(instrText)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar)

        run = paragraph.add_run("(List of tables will be updated when you open the document in Word)")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)

        run = paragraph.add_run()
        fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar)

        self.doc.add_paragraph()

    # ── Narrative ───────────────────────────────────────────────────

    def _add_narrative(self):
        """Add the introductory narrative paragraph."""
        conditions_text = ", ".join(c.label for c in sorted(self.data.conditions, key=lambda c: c.display_order))
        lots_count = len(self.data.lots)

        narrative = (
            f"This section presents the {self.data.study_type_label.lower()} stability data "
            f"for {self.data.product_name}. "
            f"A total of {lots_count} lot(s) were placed on stability under the following "
            f"storage conditions: {conditions_text}. "
            f"The results of these studies are summarized in the tables below."
        )

        p = self.doc.add_paragraph(narrative)
        p.style.font.size = self.BODY_FONT_SIZE
        self.doc.add_paragraph()

    # ── Summary Table ───────────────────────────────────────────────

    def _add_summary_table(self):
        """Add the summary table listing all lots and their study details."""
        self.table_counter += 1
        table_number = f"Table {self.data.ctd_section}-{self.table_counter}"

        # Table caption
        caption = (
            f"{table_number}: Summary of {self.data.study_type_label} "
            f"Stability Studies for {self.data.product_name}"
        )
        p = self.doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

        # Build table
        headers = [
            "Lot Number",
            "Lot Use",
            "Study Start",
            "Conditions / Timepoints Presented",
            "Table Reference",
        ]
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._style_header_cell(cell)

        # Data rows
        for row_data in self.data.summary_rows:
            row = table.add_row()
            row.cells[0].text = row_data.lot_number
            row.cells[1].text = row_data.lot_use or ""
            row.cells[2].text = row_data.study_start or ""
            row.cells[3].text = row_data.conditions_presented
            row.cells[4].text = row_data.detail_table_refs

            for cell in row.cells:
                self._style_data_cell(cell)

        self.doc.add_paragraph()

    # ── Detail Tables ───────────────────────────────────────────────

    def _add_detail_tables(self):
        """Add one detail stability table per lot × condition combination."""
        # Group by condition for section breaks
        conditions_order = sorted(self.data.conditions, key=lambda c: c.display_order)

        for condition in conditions_order:
            # Condition sub-heading
            self.doc.add_heading(f"Storage Condition: {condition.label}", level=2)

            # Filter detail tables for this condition
            tables_for_condition = [
                dt for dt in self.data.detail_tables
                if dt.condition.condition_id == condition.condition_id
            ]

            for dt in tables_for_condition:
                self._add_single_detail_table(dt)

    def _add_single_detail_table(self, dt: DetailTableData):
        """Render one detail stability table."""
        self.table_counter += 1
        table_number = f"Table {self.data.ctd_section}-{self.table_counter}"

        # Caption
        caption = (
            f"{table_number}: {self.data.study_type_label} Stability Data for "
            f"Lot {dt.lot.lot_number} at {dt.condition.label}"
        )
        p = self.doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

        # Columns: Attribute | Acceptance Criteria | TP1 | TP2 | TP3 | ...
        tp_labels = [tp.label for tp in sorted(dt.timepoints, key=lambda t: t.sort_order)]
        col_count = 2 + len(tp_labels)

        table = self.doc.add_table(rows=1, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        headers = ["Analytical Procedure / Quality Attribute", "Timepoint / Acceptance Criteria"] + tp_labels
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._style_header_cell(cell)

        # Data rows, grouped by method_group
        current_group = None
        for attr_row in dt.rows:
            # Insert group header row if method_group changes
            if attr_row.method_group and attr_row.method_group != current_group:
                current_group = attr_row.method_group
                group_row = table.add_row()
                group_row.cells[0].text = current_group
                self._style_group_header_cell(group_row.cells[0])
                # Merge remaining cells visually by leaving them empty
                for cell in group_row.cells[1:]:
                    self._style_group_header_cell(cell)

            # Data row
            row = table.add_row()
            row.cells[0].text = attr_row.attribute_name
            row.cells[1].text = attr_row.acceptance_criteria or ""

            for tp_idx, tp_label in enumerate(tp_labels):
                value = attr_row.values.get(tp_label, "–")
                row.cells[2 + tp_idx].text = str(value) if value else "–"

            for cell in row.cells:
                self._style_data_cell(cell)

        # Footnotes
        all_footnotes = dt.footnotes + self.data.global_footnotes
        if all_footnotes:
            footnote_text = "\n".join(f"  {fn}" for fn in all_footnotes)
            p = self.doc.add_paragraph()
            run = p.add_run(footnote_text)
            run.font.size = Pt(8)
            run.font.italic = True

        # Page break after each table (except possibly the last)
        self.doc.add_page_break()

    # ── Cell styling helpers ────────────────────────────────────────

    def _style_header_cell(self, cell):
        """Apply header styling: bold, background color, small font."""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = self.TABLE_FONT_SIZE
                run.font.name = self.FONT_NAME

        # Background color
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{self.HEADER_BG_COLOR}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)

    def _style_data_cell(self, cell):
        """Apply data cell styling: small font, centered."""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = self.TABLE_FONT_SIZE
                run.font.name = self.FONT_NAME

    def _style_group_header_cell(self, cell):
        """Apply method group header styling: bold, light gray bg."""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = self.TABLE_FONT_SIZE
                run.font.name = self.FONT_NAME

        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)


# ── PDF conversion helper ──────────────────────────────────────────

def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert DOCX to PDF using LibreOffice headless mode.
    Requires LibreOffice installed on the system.

    Alternative: use docx2pdf on Windows (requires MS Word).
    """
    import subprocess
    import platform

    output_dir = str(Path(pdf_path).parent)

    if platform.system() == "Windows":
        try:
            # Try docx2pdf first (requires MS Word)
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return pdf_path
        except ImportError:
            pass

    # Fall back to LibreOffice
    cmd = [
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", output_dir, docx_path,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"PDF conversion failed: {result.stderr}")

    return pdf_path

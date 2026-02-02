"""
Extraction pipeline for stability documents.

Parses PDF, DOCX, and XLSX files to extract structured stability data
(studies, lots, conditions, timepoints, assays, results) with source anchors.

This module implements a human-in-the-loop workflow: all extracted data is
stored as 'pending_review' and requires user confirmation before generation.
"""

import re
import logging
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional
from uuid import uuid4

logger = logging.getLogger(__name__)


# ── Data classes for extraction output ──────────────────────────────

@dataclass
class SourceAnchorData:
    """Links an extracted datum to its location in the source document."""
    document_id: str
    page_number: Optional[int] = None
    section_ref: Optional[str] = None
    table_ref: Optional[str] = None
    row_index: Optional[int] = None
    col_index: Optional[int] = None
    bounding_box: Optional[dict] = None
    text_snippet: Optional[str] = None


@dataclass
class ExtractedEntity:
    """Base class for all extracted entities."""
    id: str = field(default_factory=lambda: str(uuid4()))
    confidence: float = 0.0
    source_anchors: list[SourceAnchorData] = field(default_factory=list)


@dataclass
class ExtractedStudy(ExtractedEntity):
    study_type: str = ""
    study_label: str = ""
    protocol_id: Optional[str] = None
    start_date: Optional[str] = None
    sites: list[str] = field(default_factory=list)
    manufacturers: list[str] = field(default_factory=list)


@dataclass
class ExtractedLot(ExtractedEntity):
    lot_number: str = ""
    manufacturer: Optional[str] = None
    manufacturing_site: Optional[str] = None
    intended_use: Optional[str] = None
    lot_use_label: Optional[str] = None


@dataclass
class ExtractedCondition(ExtractedEntity):
    label: str = ""
    temperature_setpoint: Optional[float] = None
    tolerance: Optional[str] = None
    humidity: Optional[str] = None


@dataclass
class ExtractedTimepoint(ExtractedEntity):
    value: float = 0.0
    unit: str = "month"
    label: str = ""
    sort_order: int = 0


@dataclass
class ExtractedAttribute(ExtractedEntity):
    name: str = ""
    method_group: Optional[str] = None
    analytical_procedure: Optional[str] = None
    acceptance_criteria_text: Optional[str] = None


@dataclass
class ExtractedResult(ExtractedEntity):
    lot_ref: str = ""           # lot_number to link
    condition_ref: str = ""     # condition label to link
    timepoint_ref: str = ""     # timepoint label to link
    attribute_ref: str = ""     # attribute name to link
    value_text: Optional[str] = None
    value_numeric: Optional[float] = None
    status: Optional[str] = None
    unit: Optional[str] = None


@dataclass
class ExtractionResult:
    """Complete extraction output for one document."""
    document_id: str
    studies: list[ExtractedStudy] = field(default_factory=list)
    lots: list[ExtractedLot] = field(default_factory=list)
    conditions: list[ExtractedCondition] = field(default_factory=list)
    timepoints: list[ExtractedTimepoint] = field(default_factory=list)
    attributes: list[ExtractedAttribute] = field(default_factory=list)
    results: list[ExtractedResult] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)


# ── File type router ────────────────────────────────────────────────

class FileTypeRouter:
    """Routes documents to the appropriate extractor based on file type."""

    def route(self, file_path: str, file_type: str, document_id: str) -> "BaseExtractor":
        extractors = {
            "pdf": PDFExtractor,
            "docx": DOCXExtractor,
            "xlsx": XLSXExtractor,
        }
        ext = file_type.lower()
        if ext not in extractors:
            raise ValueError(f"Unsupported file type: {ext}")
        return extractors[ext](file_path, document_id)


# ── Base extractor ──────────────────────────────────────────────────

class BaseExtractor:
    """Base class for document extractors."""

    def __init__(self, file_path: str, document_id: str):
        self.file_path = file_path
        self.document_id = document_id

    def extract(self) -> ExtractionResult:
        raise NotImplementedError


# ── PDF extractor ───────────────────────────────────────────────────

class PDFExtractor(BaseExtractor):
    """
    Extracts stability data from PDF documents using pdfplumber.

    Strategy:
    1. Extract full text with page numbers for narrative/metadata.
    2. Extract tables using pdfplumber's table detection.
    3. Apply table template matching to identify stability table layouts.
    4. Parse entities (lots, conditions, assays, results) from matched tables.
    5. Fall back to regex/heuristic extraction for unstructured text.
    """

    # Patterns for stability table headers
    TABLE_TITLE_PATTERN = re.compile(
        r"Table\s+([\d.]+[-–]\d+)\s*[:\s]*(.+)",
        re.IGNORECASE,
    )

    CONDITION_PATTERN = re.compile(
        r"(-?\d+)\s*(?:to\s*(-?\d+))?\s*°?\s*C"
        r"|(-?\d+)\s*±\s*(\d+)\s*°?\s*C",
        re.IGNORECASE,
    )

    LOT_PATTERN = re.compile(
        r"(?:Lot|Batch)\s*(?:#|No\.?|Number)?\s*[:\s]*([A-Z0-9][-A-Z0-9]+)",
        re.IGNORECASE,
    )

    TIMEPOINT_PATTERN = re.compile(
        r"\b(\d+)\s*(W|M|D|H|week|month|day|hour)s?\b",
        re.IGNORECASE,
    )

    STUDY_TYPE_KEYWORDS = {
        "accelerated": "accelerated",
        "long-term": "long_term",
        "long term": "long_term",
        "intermediate": "intermediate",
        "stress": "stress",
        "photostability": "photostability",
    }

    def extract(self) -> ExtractionResult:
        """Run the full extraction pipeline on a PDF."""
        import pdfplumber

        result = ExtractionResult(document_id=self.document_id)

        try:
            with pdfplumber.open(self.file_path) as pdf:
                # Step 1: Extract text and identify sections
                pages_text = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    pages_text.append((i + 1, text))

                # Step 2: Extract study metadata from narrative text
                result.studies = self._extract_studies(pages_text)

                # Step 3: Extract tables
                for i, page in enumerate(pdf.pages):
                    page_num = i + 1
                    tables = page.extract_tables()
                    page_text = pages_text[i][1]

                    for table_idx, table in enumerate(tables):
                        self._process_table(
                            table, page_num, table_idx, page_text, result
                        )

                # Step 4: Extract lots from text if not found in tables
                if not result.lots:
                    result.lots = self._extract_lots_from_text(pages_text)

                # Step 5: Deduplicate
                self._deduplicate(result)

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}", exc_info=True)
            result.errors.append(f"PDF extraction error: {str(e)}")

        return result

    def _extract_studies(self, pages_text: list[tuple[int, str]]) -> list[ExtractedStudy]:
        """Extract study metadata from document text."""
        studies = []
        full_text = "\n".join(text for _, text in pages_text)

        for keyword, study_type in self.STUDY_TYPE_KEYWORDS.items():
            if keyword.lower() in full_text.lower():
                # Find the page where this appears
                page_num = None
                snippet = ""
                for pn, text in pages_text:
                    if keyword.lower() in text.lower():
                        page_num = pn
                        # Get surrounding context
                        idx = text.lower().index(keyword.lower())
                        snippet = text[max(0, idx - 100):idx + 200]
                        break

                study = ExtractedStudy(
                    study_type=study_type,
                    study_label=f"{keyword.title()} Stability Study",
                    confidence=0.7,
                    source_anchors=[SourceAnchorData(
                        document_id=self.document_id,
                        page_number=page_num,
                        text_snippet=snippet[:300],
                    )],
                )
                studies.append(study)

        return studies

    def _process_table(
        self,
        table: list[list],
        page_num: int,
        table_idx: int,
        page_text: str,
        result: ExtractionResult,
    ):
        """Process a single extracted table to identify stability data."""
        if not table or len(table) < 2:
            return

        # Check if this looks like a stability results table
        header_row = table[0]
        if not header_row:
            return

        header_text = " ".join(str(c or "") for c in header_row).lower()

        # Detect stability table by looking for timepoint columns or
        # "quality attribute" / "analytical procedure" headers
        is_stability_table = any(kw in header_text for kw in [
            "quality attribute", "analytical procedure",
            "timepoint", "acceptance criteria",
            "t0", "1w", "2w", "1m", "2m", "3m",
        ])

        if is_stability_table:
            self._parse_stability_table(table, page_num, table_idx, page_text, result)

        # Check for summary table
        is_summary = any(kw in header_text for kw in [
            "lot", "batch", "study", "condition", "table number",
        ])
        if is_summary and not is_stability_table:
            self._parse_summary_table(table, page_num, table_idx, result)

    def _parse_stability_table(
        self,
        table: list[list],
        page_num: int,
        table_idx: int,
        page_text: str,
        result: ExtractionResult,
    ):
        """Parse a stability results table into normalized entities."""
        header_row = [str(c or "").strip() for c in table[0]]

        # Identify timepoint columns (columns after the first 2 which are
        # typically "Attribute" and "Acceptance Criteria")
        timepoint_cols = {}
        for col_idx, header in enumerate(header_row):
            if col_idx < 2:
                continue
            tp = self._parse_timepoint_label(header)
            if tp:
                timepoint_cols[col_idx] = tp
                # Add to result timepoints if not duplicate
                existing = {t.label for t in result.timepoints}
                if tp.label not in existing:
                    result.timepoints.append(tp)

        # Try to extract condition from table title in page text
        table_title_match = self.TABLE_TITLE_PATTERN.search(page_text)
        condition = None
        lot_ref = None
        if table_title_match:
            title_text = table_title_match.group(2)
            cond_match = self.CONDITION_PATTERN.search(title_text)
            if cond_match:
                condition = self._parse_condition(cond_match, title_text, page_num)
                existing_labels = {c.label for c in result.conditions}
                if condition.label not in existing_labels:
                    result.conditions.append(condition)

            lot_match = self.LOT_PATTERN.search(title_text)
            if lot_match:
                lot_ref = lot_match.group(1)

        # Parse data rows
        current_group = None
        for row_idx, row in enumerate(table[1:], start=1):
            if not row or len(row) < 2:
                continue

            cell0 = str(row[0] or "").strip()
            cell1 = str(row[1] or "").strip()

            # Detect method group headers (rows where only first cell has content)
            if cell0 and not cell1 and not any(row[2:]):
                current_group = cell0
                continue

            # This is an assay/attribute row
            if cell0 or cell1:
                attr_name = cell0 or current_group or ""
                criteria_text = cell1

                if attr_name:
                    # Add attribute
                    attr = ExtractedAttribute(
                        name=attr_name,
                        method_group=current_group,
                        acceptance_criteria_text=criteria_text if criteria_text else None,
                        confidence=0.85,
                        source_anchors=[SourceAnchorData(
                            document_id=self.document_id,
                            page_number=page_num,
                            table_ref=f"table_{table_idx}",
                            row_index=row_idx,
                            text_snippet=attr_name,
                        )],
                    )
                    existing_attrs = {a.name for a in result.attributes}
                    if attr_name not in existing_attrs:
                        result.attributes.append(attr)

                    # Extract results for each timepoint column
                    for col_idx, tp in timepoint_cols.items():
                        if col_idx < len(row):
                            value = str(row[col_idx] or "").strip()
                            if value:
                                res = ExtractedResult(
                                    lot_ref=lot_ref or "",
                                    condition_ref=condition.label if condition else "",
                                    timepoint_ref=tp.label,
                                    attribute_ref=attr_name,
                                    value_text=value,
                                    value_numeric=self._try_parse_float(value),
                                    status=self._infer_status(value),
                                    confidence=0.85,
                                    source_anchors=[SourceAnchorData(
                                        document_id=self.document_id,
                                        page_number=page_num,
                                        table_ref=f"table_{table_idx}",
                                        row_index=row_idx,
                                        col_index=col_idx,
                                        text_snippet=value,
                                    )],
                                )
                                result.results.append(res)

    def _parse_summary_table(
        self,
        table: list[list],
        page_num: int,
        table_idx: int,
        result: ExtractionResult,
    ):
        """Parse a summary/overview table for lot and study information."""
        header_row = [str(c or "").strip().lower() for c in table[0]]

        # Find column indices
        lot_col = next((i for i, h in enumerate(header_row) if "lot" in h or "batch" in h), None)
        use_col = next((i for i, h in enumerate(header_row) if "use" in h or "purpose" in h), None)
        mfr_col = next((i for i, h in enumerate(header_row) if "manufactur" in h or "site" in h), None)

        for row_idx, row in enumerate(table[1:], start=1):
            if not row:
                continue
            if lot_col is not None and lot_col < len(row):
                lot_val = str(row[lot_col] or "").strip()
                if lot_val:
                    lot = ExtractedLot(
                        lot_number=lot_val,
                        intended_use=str(row[use_col] or "").strip() if use_col and use_col < len(row) else None,
                        manufacturer=str(row[mfr_col] or "").strip() if mfr_col and mfr_col < len(row) else None,
                        confidence=0.85,
                        source_anchors=[SourceAnchorData(
                            document_id=self.document_id,
                            page_number=page_num,
                            table_ref=f"table_{table_idx}",
                            row_index=row_idx,
                            text_snippet=lot_val,
                        )],
                    )
                    result.lots.append(lot)

    def _extract_lots_from_text(self, pages_text: list[tuple[int, str]]) -> list[ExtractedLot]:
        """Fallback: extract lot numbers from narrative text using regex."""
        lots = []
        seen = set()
        for page_num, text in pages_text:
            for match in self.LOT_PATTERN.finditer(text):
                lot_num = match.group(1)
                if lot_num not in seen:
                    seen.add(lot_num)
                    lots.append(ExtractedLot(
                        lot_number=lot_num,
                        confidence=0.6,
                        source_anchors=[SourceAnchorData(
                            document_id=self.document_id,
                            page_number=page_num,
                            text_snippet=match.group(0),
                        )],
                    ))
        return lots

    def _parse_timepoint_label(self, label: str) -> Optional[ExtractedTimepoint]:
        """Parse a column header into a timepoint."""
        label = label.strip()
        if not label:
            return None

        # Handle "T0" or "Initial"
        if label.upper() in ("T0", "INITIAL", "0"):
            return ExtractedTimepoint(value=0, unit="month", label="T0", sort_order=0, confidence=0.95)

        # Handle "1W", "2W", "1M", "3M", etc.
        match = self.TIMEPOINT_PATTERN.match(label)
        if match:
            val = int(match.group(1))
            unit_raw = match.group(2).upper()
            unit_map = {"W": "week", "M": "month", "D": "day", "H": "hour",
                        "WEEK": "week", "MONTH": "month", "DAY": "day", "HOUR": "hour"}
            unit = unit_map.get(unit_raw, "month")

            # Compute sort order (normalize to hours for ordering)
            multipliers = {"hour": 1, "day": 24, "week": 168, "month": 730, "year": 8760}
            sort_order = val * multipliers.get(unit, 730)

            return ExtractedTimepoint(
                value=val, unit=unit, label=label.strip(),
                sort_order=sort_order, confidence=0.9,
            )

        return None

    def _parse_condition(self, match: re.Match, title_text: str, page_num: int) -> ExtractedCondition:
        """Parse a storage condition from a regex match."""
        if match.group(1) and match.group(2):
            # Range: e.g., "-60 to -30 °C"
            t_min = float(match.group(1))
            t_max = float(match.group(2))
            label = f"{int(t_min)} to {int(t_max)} °C"
            setpoint = (t_min + t_max) / 2
            tolerance = None
        elif match.group(3) and match.group(4):
            # Tolerance: e.g., "-20 ± 5 °C"
            setpoint = float(match.group(3))
            tol = float(match.group(4))
            label = f"{int(setpoint)} ± {int(tol)} °C"
            tolerance = f"± {int(tol)} °C"
            t_min = setpoint - tol
            t_max = setpoint + tol
        else:
            label = match.group(0)
            setpoint = None
            t_min = None
            t_max = None
            tolerance = None

        return ExtractedCondition(
            label=label,
            temperature_setpoint=setpoint,
            tolerance=tolerance,
            confidence=0.9,
            source_anchors=[SourceAnchorData(
                document_id=self.document_id,
                page_number=page_num,
                text_snippet=title_text[:200],
            )],
        )

    @staticmethod
    def _try_parse_float(value: str) -> Optional[float]:
        """Attempt to parse a string as a float."""
        try:
            cleaned = value.replace(",", "").replace("%", "").strip()
            return float(cleaned)
        except (ValueError, AttributeError):
            return None

    @staticmethod
    def _infer_status(value: str) -> Optional[str]:
        """Infer a status label from a result value."""
        v = value.strip().upper()
        status_map = {
            "S": "S", "MEETS": "S", "PASS": "S", "CONFORMS": "S",
            "NS": "NS", "FAILS": "NS", "DOES NOT MEET": "NS",
            "PENDING": "Pending", "NT": "NT", "N/A": "NT",
            "-": "NT", "–": "NT",
        }
        if v in status_map:
            return status_map[v]
        # If it's a numeric value, it's a reported value (meets by implication)
        try:
            float(v.replace(",", "").replace("%", ""))
            return "S"  # assume meets if numeric
        except ValueError:
            return None

    def _deduplicate(self, result: ExtractionResult):
        """Remove duplicate entities, keeping highest confidence."""
        # Deduplicate conditions by label
        seen = {}
        deduped = []
        for c in result.conditions:
            if c.label not in seen or c.confidence > seen[c.label].confidence:
                seen[c.label] = c
        result.conditions = list(seen.values())

        # Deduplicate lots by lot_number
        seen = {}
        for lot in result.lots:
            if lot.lot_number not in seen or lot.confidence > seen[lot.lot_number].confidence:
                seen[lot.lot_number] = lot
        result.lots = list(seen.values())


# ── DOCX extractor ──────────────────────────────────────────────────

class DOCXExtractor(BaseExtractor):
    """Extracts stability data from DOCX documents using python-docx."""

    def extract(self) -> ExtractionResult:
        from docx import Document as DocxDocument

        result = ExtractionResult(document_id=self.document_id)

        try:
            doc = DocxDocument(self.file_path)

            # Extract text from paragraphs
            paragraphs_text = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs_text.append((i, para.text.strip()))

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)

                if len(rows) >= 2:
                    # Reuse PDF extractor's table parsing logic
                    pdf_ext = PDFExtractor(self.file_path, self.document_id)
                    # Get surrounding text for table title detection
                    page_text = " ".join(t for _, t in paragraphs_text)
                    pdf_ext._process_table(rows, 1, table_idx, page_text, result)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}", exc_info=True)
            result.errors.append(f"DOCX extraction error: {str(e)}")

        return result


# ── XLSX extractor ──────────────────────────────────────────────────

class XLSXExtractor(BaseExtractor):
    """Extracts stability data from XLSX spreadsheets."""

    def extract(self) -> ExtractionResult:
        import openpyxl

        result = ExtractionResult(document_id=self.document_id)

        try:
            wb = openpyxl.load_workbook(self.file_path, data_only=True)

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    rows.append([str(c) if c is not None else "" for c in row])

                if len(rows) >= 2:
                    pdf_ext = PDFExtractor(self.file_path, self.document_id)
                    pdf_ext._process_table(rows, 1, 0, sheet_name, result)

        except Exception as e:
            logger.error(f"XLSX extraction failed: {e}", exc_info=True)
            result.errors.append(f"XLSX extraction error: {str(e)}")

        return result


# ── Orchestrator ────────────────────────────────────────────────────

class ExtractionOrchestrator:
    """
    Orchestrates the extraction pipeline across multiple documents.

    Usage:
        orchestrator = ExtractionOrchestrator()
        results = orchestrator.run(documents)
        merged = orchestrator.merge(results)
    """

    def __init__(self):
        self.router = FileTypeRouter()

    def run_single(self, file_path: str, file_type: str, document_id: str) -> ExtractionResult:
        """Extract from a single document."""
        extractor = self.router.route(file_path, file_type, document_id)
        return extractor.extract()

    def run_batch(
        self,
        documents: list[dict],  # each: {file_path, file_type, document_id}
    ) -> list[ExtractionResult]:
        """Extract from multiple documents."""
        results = []
        for doc in documents:
            result = self.run_single(doc["file_path"], doc["file_type"], doc["document_id"])
            results.append(result)
        return results

    def merge(self, results: list[ExtractionResult]) -> ExtractionResult:
        """
        Merge extraction results from multiple documents into one,
        deduplicating entities by key fields.
        """
        merged = ExtractionResult(document_id="merged")

        for r in results:
            merged.studies.extend(r.studies)
            merged.lots.extend(r.lots)
            merged.conditions.extend(r.conditions)
            merged.timepoints.extend(r.timepoints)
            merged.attributes.extend(r.attributes)
            merged.results.extend(r.results)
            merged.errors.extend(r.errors)

        # Deduplicate
        merged.conditions = self._dedup_by_key(merged.conditions, lambda c: c.label)
        merged.lots = self._dedup_by_key(merged.lots, lambda l: l.lot_number)
        merged.timepoints = self._dedup_by_key(merged.timepoints, lambda t: (t.value, t.unit))
        merged.attributes = self._dedup_by_key(merged.attributes, lambda a: a.name)

        # Sort timepoints
        merged.timepoints.sort(key=lambda t: t.sort_order)

        return merged

    @staticmethod
    def _dedup_by_key(items: list, key_fn) -> list:
        """Keep highest-confidence item for each key."""
        seen = {}
        for item in items:
            k = key_fn(item)
            if k not in seen or item.confidence > seen[k].confidence:
                seen[k] = item
        return list(seen.values())

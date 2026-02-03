"""
Extract plain text from uploaded documents (PDF, DOCX, XLSX).

SERVERLESS-COMPATIBLE: Works with in-memory bytes, not filesystem paths.
Used to feed document content into the Claude generation prompt.
"""

import io
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook


def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
    """
    Extract text from file bytes based on file type.

    SERVERLESS: This function works entirely in-memory, no filesystem needed.

    Args:
        file_bytes: Raw file content as bytes
        file_type: File extension without dot (e.g., "pdf", "docx", "xlsx")

    Returns:
        Extracted text content
    """
    ext = file_type.lower().lstrip(".")

    if ext == "pdf":
        return _extract_pdf_bytes(file_bytes)
    elif ext in ("docx", "doc"):
        return _extract_docx_bytes(file_bytes)
    elif ext in ("xlsx", "xls"):
        return _extract_xlsx_bytes(file_bytes)
    else:
        return ""


def _extract_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    pages = []
    # pdfplumber can open from a file-like object
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

            # Also extract tables as text
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        cells = [str(c).strip() if c else "" for c in row]
                        pages.append(" | ".join(cells))
    return "\n\n".join(pages)


def _extract_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    # python-docx can open from a file-like object
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_xlsx_bytes(file_bytes: bytes) -> str:
    """Extract text from XLSX bytes."""
    # openpyxl can open from a file-like object
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# ── Legacy file-path API (for local development compatibility) ──────

def extract_text(file_path: str) -> str:
    """
    Extract text from a file path (legacy API for local dev).

    For serverless, use extract_text_from_bytes() instead.
    """
    path = Path(file_path)
    if not path.exists():
        return ""

    file_bytes = path.read_bytes()
    return extract_text_from_bytes(file_bytes, path.suffix)
